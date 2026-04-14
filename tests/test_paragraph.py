import unittest

from docx.enum.style import WD_STYLE_TYPE

from md2gost.layout_tracker import LayoutTracker
from md2gost.renderable.paragraph import Paragraph

from . import _create_test_document, _EMUS_PER_PX


class TestParagraph(unittest.TestCase):
    def setUp(self) -> None:
        self._document, self._max_height, self._max_width = _create_test_document()

    def test_render(self):
        paragraph = Paragraph(self._document._body)
        layout_tracker = LayoutTracker(self._max_height, self._max_width)

        paragraph.add_run("hello world")
        info = list(paragraph.render(None, layout_tracker.current_state))[0]

        self.assertAlmostEqual(45.5, info.height / _EMUS_PER_PX, delta=1/3)

    def test_add_link_anchor_explicit_none_does_not_set_hyperlink_style(self):
        paragraph = Paragraph(self._document._body)
        link = paragraph.add_link_anchor("anchor-id", None)
        link.add_run("toc item")

        styles = paragraph._docx_paragraph._p.xpath(".//w:hyperlink//w:rPr/w:rStyle/@w:val")
        self.assertEqual([], styles)

    def test_add_link_anchor_default_uses_hyperlink_style(self):
        paragraph = Paragraph(self._document._body)
        if "Hyperlink" not in [style.name for style in self._document.styles]:
            self._document.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
        link = paragraph.add_link_anchor("anchor-id")
        link.add_run("regular link")

        styles = paragraph._docx_paragraph._p.xpath(".//w:hyperlink//w:rPr/w:rStyle/@w:val")
        self.assertTrue(all(style == "Hyperlink" for style in styles))
        self.assertGreater(len(styles), 0)

