"""Security tests for image/code attachments path handling."""

import unittest
from pathlib import Path

import docx

from md2gost.extended_markdown import markdown
from md2gost.layout_tracker import LayoutTracker
from md2gost.parser_ import Parser
from md2gost.renderable.image import Image
from md2gost.renderable_factory import RenderableFactory
from md2gost.warnings_collector import clear_warnings, get_warnings

from . import _create_test_document


class TestParserPathSecurity(unittest.TestCase):
    def test_blocks_external_image_reference(self):
        parsed = markdown.parse("![img](https://example.com/image.png)\n")
        paragraph = parsed.children[0]

        Parser.resolve_paths(paragraph, "/tmp/workdir")

        image = paragraph.children[0]
        self.assertEqual(getattr(image, "source_dest"), "https://example.com/image.png")
        self.assertEqual(image.dest, "__blocked_external_reference__")

    def test_blocks_image_path_traversal(self):
        parsed = markdown.parse("![img](../../etc/passwd)\n")
        paragraph = parsed.children[0]

        Parser.resolve_paths(paragraph, "/tmp/workdir")

        image = paragraph.children[0]
        self.assertEqual(getattr(image, "source_dest"), "../../etc/passwd")
        self.assertEqual(image.dest, "__blocked_path_traversal__")

    def test_blocks_external_code_attachment(self):
        parsed = markdown.parse("```python https://example.com/code.py\nprint('ok')\n```\n")
        code_block = parsed.children[0]

        Parser.resolve_paths(code_block, "/tmp/workdir")

        self.assertEqual(getattr(code_block, "source_extra"), "https://example.com/code.py")
        self.assertEqual(code_block.extra, "__blocked_external_reference__")

    def test_blocks_code_path_traversal(self):
        parsed = markdown.parse("```python ../../secrets.py\nprint('ok')\n```\n")
        code_block = parsed.children[0]

        Parser.resolve_paths(code_block, "/tmp/workdir")

        self.assertEqual(getattr(code_block, "source_extra"), "../../secrets.py")
        self.assertEqual(code_block.extra, "__blocked_path_traversal__")

    def test_blocks_absolute_image_path_outside_workspace(self):
        parsed = markdown.parse("![img](/tmp/tmpbzwqhz78/media/media/image2.png)\n")
        paragraph = parsed.children[0]

        Parser.resolve_paths(paragraph, "/tmp/workdir")

        image = paragraph.children[0]
        self.assertEqual(getattr(image, "source_dest"), "/tmp/tmpbzwqhz78/media/media/image2.png")
        self.assertEqual(image.dest, "__blocked_path_traversal__")


class TestAttachmentWarnings(unittest.TestCase):
    def setUp(self):
        clear_warnings()

    def test_image_external_reference_produces_warning(self):
        document, max_height, max_width = _create_test_document()
        image = Image(
            document._body,
            "__blocked_external_reference__",
            source_path="https://example.com/image.png",
        )

        rendered = list(image.render(None, LayoutTracker(max_height, max_width).current_state))

        self.assertEqual(rendered, [])
        self.assertTrue(any("Внешние URL картинок запрещены" in msg for msg in get_warnings()))

    def test_factory_warns_for_external_code_attachment(self):
        template_path = Path(__file__).resolve().parents[1] / "md2gost" / "Template.docx"
        template_doc = docx.Document(str(template_path))
        factory = RenderableFactory(template_doc._body)
        parsed = markdown.parse("```python https://example.com/code.py\nprint('ok')\n```\n")
        code_block = parsed.children[0]
        Parser.resolve_paths(code_block, "/tmp/workdir")

        list(factory.create(code_block, None))

        self.assertTrue(any("Внешние источники кода запрещены" in msg for msg in get_warnings()))

    def test_image_warning_does_not_expose_absolute_path(self):
        document, max_height, max_width = _create_test_document()
        image = Image(
            document._body,
            "__blocked_path_traversal__",
            source_path="/tmp/tmpbzwqhz78/media/media/image2.png",
        )

        rendered = list(image.render(None, LayoutTracker(max_height, max_width).current_state))

        self.assertEqual(rendered, [])
        warnings = get_warnings()
        self.assertTrue(any("media/image2.png" in msg for msg in warnings))
        self.assertTrue(all("/tmp/" not in msg for msg in warnings))

    def test_code_warning_does_not_expose_absolute_path(self):
        template_path = Path(__file__).resolve().parents[1] / "md2gost" / "Template.docx"
        template_doc = docx.Document(str(template_path))
        factory = RenderableFactory(template_doc._body)
        parsed = markdown.parse("```python /tmp/tmpbzwqhz78/media/media/code.py\nprint('ok')\n```\n")
        code_block = parsed.children[0]
        Parser.resolve_paths(code_block, "/tmp/workdir")

        list(factory.create(code_block, None))

        warnings = get_warnings()
        self.assertTrue(any("Доступ к файлам вне рабочей директории запрещён" in msg for msg in warnings))
        self.assertTrue(any("media/code.py" in msg for msg in warnings))
        self.assertTrue(all("/tmp/" not in msg for msg in warnings))

    def test_image_file_url_warning_does_not_expose_absolute_path(self):
        document, max_height, max_width = _create_test_document()
        image = Image(
            document._body,
            "__blocked_external_reference__",
            source_path="file:///tmp/md2gost_lmosisgn/Aspose.Words.x.png",
        )

        rendered = list(image.render(None, LayoutTracker(max_height, max_width).current_state))

        self.assertEqual(rendered, [])
        warnings = get_warnings()
        self.assertTrue(any("Aspose.Words.x.png" in msg for msg in warnings))
        self.assertTrue(all("/tmp/" not in msg for msg in warnings))

    def test_code_file_url_warning_does_not_expose_absolute_path(self):
        template_path = Path(__file__).resolve().parents[1] / "md2gost" / "Template.docx"
        template_doc = docx.Document(str(template_path))
        factory = RenderableFactory(template_doc._body)
        parsed = markdown.parse("```python file:///tmp/md2gost_lmosisgn/private.py\nprint('ok')\n```\n")
        code_block = parsed.children[0]
        Parser.resolve_paths(code_block, "/tmp/workdir")

        list(factory.create(code_block, None))

        warnings = get_warnings()
        self.assertTrue(any("private.py" in msg for msg in warnings))
        self.assertTrue(all("/tmp/" not in msg for msg in warnings))
