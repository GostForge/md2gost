"""Тесты для NumberingPreProcessor."""

import unittest

import docx
from docx.shared import Cm, Pt
from docx.enum.style import WD_STYLE_TYPE

from md2gost.numberer import NumberingPreProcessor
from md2gost.renderable.paragraph import Paragraph
from md2gost.renderable.heading import Heading
from md2gost.renderable.caption import CaptionInfo
from md2gost.renderable.requires_numbering import RequiresNumbering


class FakeNumbered(RequiresNumbering):
    """Минимальный RequiresNumbering для тестов."""
    def __init__(self, category: str, unique_name: str = None):
        super().__init__(category, unique_name)
        self.number = None

    def set_number(self, number: int | str):
        self.number = number

    def render(self, *args):
        yield from []


class TestNumberingPreProcessor(unittest.TestCase):
    def test_sequential_numbering(self):
        numberer = NumberingPreProcessor()
        items = [FakeNumbered("Таблица"), FakeNumbered("Таблица"), FakeNumbered("Таблица")]
        numberer.process(items)
        self.assertEqual([i.number for i in items], ["1", "2", "3"])

    def test_separate_categories(self):
        numberer = NumberingPreProcessor()
        items = [
            FakeNumbered("Таблица"), FakeNumbered("Рисунок"),
            FakeNumbered("Таблица"), FakeNumbered("Рисунок"),
        ]
        numberer.process(items)
        self.assertEqual(items[0].number, "1")  # Таблица 1
        self.assertEqual(items[1].number, "1")  # Рисунок 1
        self.assertEqual(items[2].number, "2")  # Таблица 2
        self.assertEqual(items[3].number, "2")  # Рисунок 2

    def test_unique_names_stored(self):
        numberer = NumberingPreProcessor()
        items = [
            FakeNumbered("Таблица", "goods"),
            FakeNumbered("Таблица", "prices"),
        ]
        numberer.process(items)
        self.assertEqual(numberer._reference_data["goods"], "1")
        self.assertEqual(numberer._reference_data["prices"], "2")

    def test_reference_resolves(self):
        doc = docx.Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(14)

        numberer = NumberingPreProcessor()
        paragraph = Paragraph(doc._body)
        paragraph.add_reference("goods")

        items = [
            FakeNumbered("Таблица", "goods"),
            paragraph,
        ]
        numberer.process(items)
        self.assertEqual(paragraph.references[0]._element.xpath("w:t")[0].text, "1")

    def test_duplicate_unique_name_warns(self):
        """Дублирование unique_name — не падает, а пишет warning."""
        numberer = NumberingPreProcessor()
        items = [
            FakeNumbered("Таблица", "dup"),
            FakeNumbered("Таблица", "dup"),
        ]
        # Не должно упасть
        numberer.process(items)
        self.assertEqual(items[0].number, "1")
        self.assertEqual(items[1].number, "2")

    def test_missing_reference_warns(self):
        """Ссылка на несуществующий unique_name — не падает."""
        doc = docx.Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(14)

        numberer = NumberingPreProcessor()
        paragraph = Paragraph(doc._body)
        paragraph.add_reference("nonexistent")
        numberer.process([paragraph])
        # Просто не падает

    def test_empty_renderables(self):
        numberer = NumberingPreProcessor()
        numberer.process([])  # Не падает

    def test_sectional_numbering(self):
        """Посекционная нумерация: Рисунок 1.1, 1.2, затем 2.1."""
        doc = docx.Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(14)
        doc.styles["Normal"].paragraph_format.first_line_indent = Cm(1.25)
        doc.styles["Normal"].paragraph_format.line_spacing = 1.5

        numberer = NumberingPreProcessor(sectional=True)
        h1 = Heading(doc._body, 1, numbered=True)
        h1._docx_paragraph.add_run("Глава 1")
        h2 = Heading(doc._body, 1, numbered=True)
        h2._docx_paragraph.add_run("Глава 2")

        items = [
            h1,
            FakeNumbered("Рисунок", "img1"),
            FakeNumbered("Рисунок", "img2"),
            h2,
            FakeNumbered("Рисунок", "img3"),
        ]
        numberer.process(items)
        self.assertEqual(items[1].number, "1.1")
        self.assertEqual(items[2].number, "1.2")
        self.assertEqual(items[4].number, "2.1")

    def test_sectional_numbering_references(self):
        """Ссылки при посекционной нумерации."""
        doc = docx.Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(14)
        doc.styles["Normal"].paragraph_format.first_line_indent = Cm(1.25)
        doc.styles["Normal"].paragraph_format.line_spacing = 1.5

        numberer = NumberingPreProcessor(sectional=True)
        h1 = Heading(doc._body, 1, numbered=True)
        h1._docx_paragraph.add_run("Раздел")

        paragraph = Paragraph(doc._body)
        paragraph.add_reference("tbl")

        items = [
            h1,
            FakeNumbered("Таблица", "tbl"),
            paragraph,
        ]
        numberer.process(items)
        self.assertEqual(paragraph.references[0]._element.xpath("w:t")[0].text, "1.1")
