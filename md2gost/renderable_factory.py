import logging
import os
from functools import singledispatchmethod
from os import environ
from typing import Generator

from docx.shared import Parented, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .renderable import (
    CaptionInfo,
    Equation,
    Heading,
    Image,
    Link,
    List,
    Listing,
    Paragraph,
    Renderable,
    Table,
    ToC,
)
from .config import Md2GostConfig, get_default_config
from . import extended_markdown
from .warnings_collector import add_warning


_BLOCKED_PATH_TRAVERSAL = "__blocked_path_traversal__"
_BLOCKED_EXTERNAL_REFERENCE = "__blocked_external_reference__"


def _public_path_for_warning(path: str) -> str:
    if not isinstance(path, str):
        path = str(path)

    normalized = path.replace("\\", "/")
    if "://" in normalized or normalized.startswith("//"):
        return normalized

    if os.path.isabs(normalized):
        media_index = normalized.lower().find("/media/")
        if media_index != -1:
            normalized = normalized[media_index + 1:]
        else:
            normalized = os.path.basename(normalized) or normalized

    normalized = normalized.lstrip("/").lstrip("./")
    while normalized.lower().startswith("media/media/"):
        normalized = normalized[len("media/"):]
    return normalized


class RenderableFactory:
    def __init__(self, parent: Parented, config: Md2GostConfig | None = None):
        self._parent = parent
        self._config = config or get_default_config()

    @singledispatchmethod
    def create(self, marko_element: extended_markdown.BlockElement,
               caption_info: CaptionInfo) -> Generator[Renderable, None, None]:
        msg = f"{marko_element.get_type()} не поддерживается"
        paragraph = Paragraph(self._parent)
        paragraph.add_run(msg, color=RGBColor.from_string('ff0000'))
        logging.warning(msg)
        add_warning(msg)
        yield paragraph

    @staticmethod
    def _create_runs(paragraph_or_link: Paragraph | Link, children, classes: list[type] = None):
        if not classes:
            classes = []
        for child in children:
            if isinstance(child, (extended_markdown.RawText, extended_markdown.Literal)):
                paragraph_or_link.add_run(child.children,
                                          is_bold=extended_markdown.StrongEmphasis in classes or None,
                                          is_italic=extended_markdown.Emphasis in classes or None,
                                          strike_through=extended_markdown.Strikethrough in classes or None)
            elif isinstance(child, extended_markdown.CodeSpan):
                paragraph_or_link.add_run(child.children, is_italic=True)
            elif isinstance(child, (extended_markdown.LineBreak, extended_markdown.Image)):
                paragraph_or_link.add_run(" ")
            elif isinstance(child, extended_markdown.Reference):
                paragraph_or_link.add_reference(child.unique_name)
            elif isinstance(child, extended_markdown.InlineEquation):
                paragraph_or_link.add_inline_equation(child.latex_equation)
            elif isinstance(child, (extended_markdown.Link, extended_markdown.Url)):
                RenderableFactory._create_runs(paragraph_or_link.add_link_url(child.dest),
                                               child.children, classes)
            elif isinstance(child, (extended_markdown.Emphasis, extended_markdown.StrongEmphasis,
                                    extended_markdown.Strikethrough)):
                RenderableFactory._create_runs(paragraph_or_link,
                                               child.children, classes + [type(child)])
            else:
                msg = f"{child.get_type()} не поддерживается"
                paragraph_or_link.add_run(f" {msg} ",
                                          color=RGBColor.from_string("FF0000"))
                logging.warning(msg)
                add_warning(msg)

    @create.register
    def _(self, marko_paragraph: extended_markdown.Paragraph, caption_info: CaptionInfo):
        paragraph = Paragraph(self._parent)

        all_images = True
        for child in marko_paragraph.children:
            if isinstance(child, extended_markdown.Image):
                yield Image(
                    self._parent,
                    child.dest,
                    CaptionInfo(child.unique_name, child.title),
                    source_path=getattr(child, "source_dest", child.dest),
                    config=self._config,
                )
            else:
                all_images = False

        if all_images:
            return

        RenderableFactory._create_runs(paragraph, marko_paragraph.children)
        yield paragraph

    @create.register
    def _(self, marko_heading: extended_markdown.Heading | extended_markdown.SetextHeading, caption_info: CaptionInfo):
        heading = Heading(self._parent, marko_heading.level, marko_heading.numbered)
        RenderableFactory._create_runs(heading, marko_heading.children)
        yield heading

    @create.register
    def _(self, marko_code_block: extended_markdown.FencedCode | extended_markdown.CodeBlock, caption_info: CaptionInfo):
        listing = Listing(self._parent, marko_code_block.lang, caption_info, config=self._config)

        text = marko_code_block.children[0].children
        source_extra = _public_path_for_warning(getattr(marko_code_block, "source_extra", marko_code_block.extra))
        if marko_code_block.extra:
            if marko_code_block.extra == _BLOCKED_EXTERNAL_REFERENCE:
                msg = f"Внешние источники кода запрещены политикой безопасности: {source_extra}"
                logging.warning(msg)
                add_warning(msg)
            elif marko_code_block.extra == _BLOCKED_PATH_TRAVERSAL:
                msg = f"Доступ к файлам вне рабочей директории запрещён: {source_extra}"
                logging.warning(msg)
                add_warning(msg)
            else:
                try:
                    with open(marko_code_block.extra, encoding="utf-8") as f:
                        text = f.read() + text
                except FileNotFoundError:
                    msg = f"Файл с кодом не найден: {source_extra}"
                    logging.warning(msg)
                    add_warning(msg)

        listing.set_text(text)
        yield listing

    @create.register
    def _(self, marko_thematic_break: extended_markdown.ThematicBreak, caption_info: CaptionInfo):
        yield from []

    @create.register
    def _(self, marko_equation: extended_markdown.Equation, caption_info: CaptionInfo):
        equation = Equation(self._parent, marko_equation.latex_equation, caption_info, config=self._config)
        yield equation

    @create.register
    def _(self, marko_list: extended_markdown.List, caption_info: CaptionInfo):
        list_ = List(self._parent, marko_list.ordered, config=self._config)

        def create_items_from_marko(marko_list_, level=1):
            for list_item in marko_list_.children:
                for child in list_item.children:
                    if isinstance(child, extended_markdown.List):
                        create_items_from_marko(child, level + 1)
                    elif isinstance(child, extended_markdown.Paragraph):
                        RenderableFactory._create_runs(
                            list_.add_item(level),
                            child.children
                        )

        create_items_from_marko(marko_list)

        yield list_

    @create.register
    def _(self, marko_table: extended_markdown.Table, caption_info: CaptionInfo):
        table = Table(self._parent, len(marko_table.children), len(marko_table.children[0].children),
                      caption_info, config=self._config)
        for i, row in enumerate(marko_table.children):
            for j, cell in enumerate(row.children):
                paragraph = table.add_paragraph_to_cell(i, j)
                paragraph.alignment = {
                    None: None,
                    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    "center": WD_PARAGRAPH_ALIGNMENT.CENTER
                }[cell.align]
                RenderableFactory._create_runs(
                    paragraph,
                    cell.children
                )

        yield table

    @create.register
    def _(self, marko_toc: extended_markdown.TOC, caption_info: CaptionInfo):
        toc = ToC(self._parent)
        yield toc
